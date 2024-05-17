import os
# from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
import docx
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_history_aware_retriever
from langchain_core.prompts import MessagesPlaceholder



# Function to read Word documents
def read_docx(file_path):
    doc = docx.Document(file_path)
    content = [Document(page_content=paragraph.text) for paragraph in doc.paragraphs]
    return content

# Function to read and process various document formats
def read_document(folder):  
    combined_document= []
    # Read and combine content from all documents
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        if file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
            document = loader.load_and_split()
            combined_document += document
        elif file_path.endswith('.docx'):
            document = read_docx(file_path)
            combined_document += document
        elif file_path.endswith('.txt'):
            txtloader = TextLoader(file_path)
            document = txtloader.load_and_split()
            combined_document += document
    return combined_document

def check_for_faiss_file(directory):
    # List all files in the given directory
    for filename in os.listdir(directory):
        # Check if the file ends with .faiss
        if filename.endswith('.faiss'):
            return True  # A file ending with .faiss was found
    return False  # No file ending with .faiss was found

def vector_store(target_dir):
    embeddings = OpenAIEmbeddings()
    if check_for_faiss_file(target_dir):
        vector = FAISS.load_local(target_dir, embeddings = embeddings,  allow_dangerous_deserialization=True)
        return vector
    document = read_document(target_dir)
    text_splitter = RecursiveCharacterTextSplitter()
    document = text_splitter.split_documents(document)
    vector = FAISS.from_documents(document, embeddings)
    vector.save_local(target_dir)
    return vector

def build_retriever_chain(llm, retriever):
    # First we need a prompt that we can pass into an LLM to generate this search query
    prompt = ChatPromptTemplate.from_messages([
        MessagesPlaceholder(variable_name="chat_history"),
        ("user", "{input}"),
        ("user", "Given the above conversation, generate a search query to look up to get information relevant to the conversation")
    ])
    retriever_chain = create_history_aware_retriever(llm, retriever, prompt)

    return retriever_chain
